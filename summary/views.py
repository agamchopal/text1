import nltk
nltk.download("all")

nltk.download('popular')


# Download the stopwords and punkt tokenizer
nltk.download('stopwords')
nltk.download('punkt')
from django.shortcuts import render
from requests import request
from .utils import  robin
import re
import nltk
from .utils import download_summary_as_docx




def calculate_summary_percentage(summary, text):
    if 'Summary' in summary:
        summary_text = summary['Summary']  
        summary_words = summary_text.split()  
    else:
        summary_words = []

    text_words = text.split() 

    
    if len(text_words) > 0:
        summary_percentage =int( (len(summary_words) / len(text_words)) * 100)
    else:
        summary_percentage = 0

    return summary_percentage
def calculate_auto_summary_percentage(auto_summary, text):
    auto_summary_words = auto_summary.split()  

    text_words = text.split() 

   
    if len(text_words) > 0:
        auto_summary_percentage = int((len(auto_summary_words) / len(text_words)) * 100)
    else:
        auto_summary_percentage = 0

    return auto_summary_percentage


def remove_non_ascii(text):
    clean_text = re.sub(r'[^\w\s.]', '', text)
    return clean_text

import re
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from collections import Counter

def preprocess_text(text):
    
    text = re.sub(r"'s\b", "", text)
    
   
    text = re.sub(r's\b', '', text)

    return text

def lemmatize_words(words):
    lemmatizer = WordNetLemmatizer()
    return [lemmatizer.lemmatize(word) for word in words]
def summarize_and_highlight(text, min_keyword_length=3, min_keyword_count=5):
    def preprocess_text(text):
        # Add your preprocessing code here
        pass

    def lemmatize_words(words):
        # Add your lemmatization code here
        pass

    

def summarize_text(text, min_keyword_length=3, min_keyword_count=5):
    text = preprocess_text(text)
    
    words = word_tokenize(text)
    stop_words = set(stopwords.words('english'))
    punctuation = set(string.punctuation)

    
    filtered_words = [word.lower() for word in words if word.lower() not in stop_words and word not in punctuation]
    lemmatized_words = lemmatize_words(filtered_words)

    keyword_counts = Counter(lemmatized_words)

    top_keywords = keyword_counts.most_common(10)

    filtered_top_keywords = [(keyword, count) for keyword, count in top_keywords if len(keyword) >= min_keyword_length and count >= min_keyword_count]

    return filtered_top_keywords





'''def tool(request):
    context = {}
    auto_summary = None
    auto_generated_summary_word_count = None
    auto_generated_summary_char_count = None
    auto_generated_summary_char_count_one = None
    auto_generated_summary_page_count = None
    auto_generated_summary_lines = None
    auto_generated_summary_para = None
    summary_percentage = None
    


    if request.method == 'POST':
        show_statistics = False
        text = request.POST.get("description")
        keywords = request.POST.get("keywords")
        summary_length = request.POST.get("summary_length")

        context['oldDescription'] = text

        if 'extract_keywords' in request.POST:
            cleaned_text = remove_non_ascii(text)
            top_keywords = summarize_text(cleaned_text)
            context['top_keywords'] = top_keywords

        if 'generate_summary' in request.POST:
            cleaned_text = remove_non_ascii(text)
            top_keywords = summarize_text(cleaned_text)
            context['top_keywords'] = top_keywords
            summary = robin(keywords, text, summary_length)
            context['context'] = summary
            context['show_statistics'] = True
            summary_percentage = calculate_summary_percentage(summary, text)
            context['summary_percentage'] = summary_percentage

        if 'auto_generate_summary' in request.POST:
            auto_summary = auto_generate_summary(text)
            #auto_summary = summarize_and_highlight(text)
            #auto_summary = generated summary
            

            auto_generated_summary_word_count = count_words(auto_summary)
            auto_generated_summary_char_count = character_count(auto_summary)
            auto_generated_summary_char_count_one = count_char_one(auto_summary)
            auto_generated_summary_page_count = count_pages(auto_summary)
            auto_generated_summary_lines = count_lines_helper(auto_summary)
            auto_generated_summary_para = count_paragraphs_helper(auto_summary)
            summary_percentage = calculate_auto_summary_percentage(auto_summary, text)
            context['summary_percentage'] = summary_percentage

        count = count_paragraphs_helper(text)
        lines = count_lines_helper(text)
        updated_lines = count_lines(text)
        char_count = character_count(text)
        char_count_one = count_char_one(text)
        word_count = count_words(text)
        pages = count_pages(text)

        show_statistics = True

        context.update({
            'count': count,
            'lines': lines,
            'pages': pages,
            'text': text,
            'updated_lines': updated_lines,
            'char_count': char_count,
            'char_count_one': char_count_one,
            'word_count': word_count,
            'show_statistics': show_statistics,
            'keywords': keywords,
            'auto_generated_summary': auto_summary,
            'auto_generated_summary_word_count': auto_generated_summary_word_count,
            'auto_generated_summary_char_count': auto_generated_summary_char_count,
            'auto_generated_summary_char_count_one': auto_generated_summary_char_count_one,
            'auto_generated_summary_page_count': auto_generated_summary_page_count,
            'auto_generated_summary_lines': auto_generated_summary_lines,
            'auto_generated_summary_para': auto_generated_summary_para,



        })

    return render(request, 'textsum.html', context)
'''
def tool(request):
    context = {}
    auto_summary = None
    auto_generated_summary_word_count = None
    auto_generated_summary_char_count = None
    auto_generated_summary_char_count_one = None
    auto_generated_summary_page_count = None
    auto_generated_summary_lines = None
    auto_generated_summary_para = None
    summary_percentage = None
    


    if request.method == 'POST':
        show_statistics = False
        text = request.POST.get("description")
        keywords = request.POST.get("keywords")
        summary_length = request.POST.get("summary_length")

        context['oldDescription'] = text

        if 'extract_keywords' in request.POST:
            cleaned_text = remove_non_ascii(text)
            top_keywords = summarize_text(cleaned_text)
            context['top_keywords'] = top_keywords

        if 'generate_summary' in request.POST:
            cleaned_text = remove_non_ascii(text)
            top_keywords = summarize_text(cleaned_text)
            context['top_keywords'] = top_keywords
            summary = robin(keywords, text, summary_length)
            context['context'] = summary
            context['show_statistics'] = True

            
            summary_percentage = calculate_summary_percentage(summary, text)
            context['summary_percentage'] = summary_percentage
            context['Summary'] = summary


        if 'download_summary' in request.POST:
            summary = context.get('Summary', '')  # Assuming context['Summary'] contains the summary text

            # Generate DOCX file
            doc = Document()
            doc.add_paragraph(summary)
            
            # In-memory output stream for the document
            f = io.BytesIO()
            doc.save(f)
            length = f.tell()
            f.seek(0)

            response = HttpResponse(f.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename="summary.docx"'
            response['Content-Length'] = length
            return response

        if 'auto_summary' in request.POST:
            auto_summary =auto_generate_summary(text)
            
        

            auto_generated_summary_word_count = count_words(auto_summary)
            auto_generated_summary_char_count = character_count(auto_summary)
            auto_generated_summary_char_count_one = count_char_one(auto_summary)
            auto_generated_summary_page_count = count_pages(auto_summary)
            auto_generated_summary_lines = count_lines_helper(auto_summary)
            auto_generated_summary_para = count_paragraphs_helper(auto_summary)
            summary_percentage = calculate_auto_summary_percentage(auto_summary, text)
            context['summary_percentage'] = summary_percentage

        if 'auto_generate_summary2' in request.POST:
            auto_summary = robin(keywords,text, summary_length) # Assuming auto_summary contains the summary text
            # Generate DOCX file
            doc = Document()
            doc.add_paragraph(summary)
            
            # In-memory output stream for the document
            f = io.BytesIO()
            doc.save(f)
            length = f.tell()
            f.seek(0)

            response = HttpResponse(f.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename="summary.docx"'
            response['Content-Length'] = length
            return response
        

        if 'auto_generate_summary' in request.POST:
            auto_summary = auto_generate_summary(text)  # Assuming auto_summary contains the summary text
            
            # Generate DOCX file
            doc = Document()
            doc.add_paragraph(auto_summary)
            
            # In-memory output stream for the document
            f = io.BytesIO()
            doc.save(f)
            length = f.tell()
            f.seek(0)

            response = HttpResponse(f.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename="auto_generated_summary.docx"'
            response['Content-Length'] = length
            return response
        
    
        count = count_paragraphs_helper(text)
        lines = count_lines_helper(text)
        updated_lines = count_lines(text)
        char_count = character_count(text)
        char_count_one = count_char_one(text)
        word_count = count_words(text)
        pages = count_pages(text)

        show_statistics = True

        context.update({
            'count': count,
            'lines': lines,
            'pages': pages,
            'text': text,
            'updated_lines': updated_lines,
            'char_count': char_count,
            'char_count_one': char_count_one,
            'word_count': word_count,
            'show_statistics': show_statistics,
            'keywords': keywords,
            'auto_generated_summary': auto_summary,
            'auto_generated_summary_word_count': auto_generated_summary_word_count,
            'auto_generated_summary_char_count': auto_generated_summary_char_count,
            'auto_generated_summary_char_count_one': auto_generated_summary_char_count_one,
            'auto_generated_summary_page_count': auto_generated_summary_page_count,
            'auto_generated_summary_lines': auto_generated_summary_lines,
            'auto_generated_summary_para': auto_generated_summary_para,



        })

    return render(request, 'textsum.html', context)


def download_summary(request):
    if request.method == 'POST':
        keywords = request.POST.get('keywords')  # Assuming you're using a form to submit keywords
        text = request.POST.get('description')          # Assuming you're using a form to submit text
        summary_length = request.POST.get('summary_length')  # Assuming you're using a form to specify summary length

        # Generate summary using the 'robin' function
        summary = robin(keywords, text, summary_length)

        # Create a new Document object
        doc = Document()

        # Add a title to the document
        doc.add_heading('Summary', level=1)

        # Add the summary to the document
        doc.add_paragraph(summary)

        # Create an in-memory output stream for the document
        f = io.BytesIO()

        # Save the document to the in-memory stream
        doc.save(f)
        length = f.tell()
        f.seek(0)

        # Create an HttpResponse with the content type set to 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response = HttpResponse(f.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        
        # Set the Content-Disposition header to indicate the filename
        response['Content-Disposition'] = 'attachment; filename="generated_summary.docx"'
        
        # Set the Content-Length header
        response['Content-Length'] = length

        return response
    else:
        return HttpResponse("Invalid Request")


def count_paragraphs_helper(text):
    paragraphs = re.split(r'\n\s*\n', text.strip())
    return len(paragraphs)

def count_lines_helper(text):
    sentences = re.split(r'\.\s+', text)

    line_count = len(sentences)
    return line_count

def count_pages(text):

    lines_per_page = 50
    lines = count_lines_helper(text)
    pages = lines // lines_per_page
    if lines % lines_per_page != 0:
        pages += 1
    return pages

def character_count(text):
    c = 0
    for i in range(len(text)):
        if text[i] != " ":
            c += 1
    return c

def count_char_one(text):
    return len(text)

def count_words(text):
    words = text.split()
    return len(words)

def count_lines(text):
    lines = re.split(r'\.\s+', text)
    formatted_lines = '<br>'.join(lines)
    return formatted_lines



import string
from django.shortcuts import render
from collections import Counter
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize


import io
import PyPDF2


from PyPDF2 import PdfReader

from django.shortcuts import render
from .utils import robin
from datetime import datetime


def format_pdf_date(pdf_date_string):
    year = int(pdf_date_string[2:6])
    month = int(pdf_date_string[6:8])
    day = int(pdf_date_string[8:10])
    hour = int(pdf_date_string[10:12])
    minute = int(pdf_date_string[12:14])
    second = int(pdf_date_string[14:16])
    return f"{year}-{month:02d}-{day:02d} {hour:02d}:{minute:02d}:{second:02d}"



def remove_non_ascii(text):
    clean_text = re.sub(r'[^\w\s]', '', text)
    return clean_text


def remove_urls(text):
    url_pattern = re.compile(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+')
    cleaned_text = re.sub(url_pattern, '', text)
    return cleaned_text





def pdf_date_to_readable(pdf_date):
    try:
 
        year = int(pdf_date[2:6])
        month = int(pdf_date[6:8])
        day = int(pdf_date[8:10])
        hour = int(pdf_date[10:12])
        minute = int(pdf_date[12:14])
        second = int(pdf_date[14:16])

   
        dt = datetime(year, month, day, hour, minute, second)

        return dt.strftime('%Y-%m-%d %H:%M:%S')
    except:
        return "Unknown"




from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
import string
from collections import Counter





from django.shortcuts import render
from .utils import robin

import io
import PyPDF2
import re

from django.contrib.sessions.models import Session


def process_pdf_view(request):
    context = {}

    if request.method == 'POST':
        if 'pdf_file' in request.FILES:
            uploaded_file = request.FILES['pdf_file']

            if not uploaded_file:
                context['error_message'] = 'The uploaded PDF file is empty.'
                return render(request, 'pdf2.html', context)

            file_contents = uploaded_file.read()

            if uploaded_file.name.lower().endswith('.pdf'):
                pdf_stream = io.BytesIO(file_contents)
                pdf_reader = PyPDF2.PdfReader(pdf_stream)

                context['uploaded_pdf'] = True
                metadata = pdf_reader.metadata
                context['title'] = metadata.get('/Title', 'Unknown')
                context['author'] = metadata.get('/Author', 'Unknown')
                context['creator'] = metadata.get('/Creator', 'Unknown')
                context['producer'] = metadata.get('/Producer', 'Unknown')

                creation_date = metadata.get('/CreationDate', '')
                context['readable_creation_date'] = pdf_date_to_readable(creation_date)

                mod_date = metadata.get('/ModDate', '')
                context['readable_mod_date'] = pdf_date_to_readable(mod_date)

                text = ""
                for page_num in range(len(pdf_reader.pages)):
                    text += pdf_reader.pages[page_num].extract_text()

                cleaned_text_without_urls = remove_urls(text)
                extracted_text = re.sub(r'\s+', ' ', cleaned_text_without_urls)

                cleaned_text = remove_non_ascii(text)
                cleaned_text_without_urls = remove_urls(cleaned_text)
                top_keywords = summarize_text(cleaned_text_without_urls)
                num_paragraphs = count_paragraphs_helper(extracted_text)
                lines = count_lines_helper(text)
                char_count = character_count(text)
                char_count_spaces = len(text)
                word_count = count_words(text)
                total_pages = len(pdf_reader.pages)


                context['top_keywords'] = top_keywords
                context['extracted_text'] = extracted_text
                context['num_paragraphs'] = num_paragraphs
                context['lines'] = lines
                context['char_count'] = char_count
                context['char_count_spaces'] = char_count_spaces
                context['word_count'] = word_count
                context['total_pages'] = total_pages

                
                request.session['extracted_data'] = {
                    'top_keywords': top_keywords,
                    'extracted_text': extracted_text,
                    'num_paragraphs': num_paragraphs,
                    'lines': lines,
                    'char_count': char_count,
                    'char_count_spaces': char_count_spaces,
                    'word_count': word_count,
                    'total_pages': total_pages,
                    'title': context['title'],
                    'author': context['author'],
                    'creator': context['creator'],
                    'producer': context['producer'],
                    'readable_creation_date': context['readable_creation_date'],
                    'readable_mod_date': context['readable_mod_date'],
                }
                
                if 'auto_generate_summary1' in request.POST:
                    auto_summary = auto_generate_summary(text)  # Assuming auto_summary contains the summary text

                
                
            # Generate DOCX file
                    doc = Document()
                    doc.add_paragraph(auto_summary)
            
            # In-memory output stream for the document
                    f = io.BytesIO()
                    doc.save(f)
                    length = f.tell()
                    f.seek(0)

                    response = HttpResponse(f.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = 'attachment; filename="auto_generated_summary.docx"'
                    response['Content-Length'] = length
                    return response


              
                
                   
                
                
                if 'auto_generate_summary' in request.POST:

                    extracted_data = request.session.get('extracted_data', {})
                    extracted_text = extracted_data.get('extracted_text', "")
                    top_keywords = extracted_data.get('top_keywords', [])
                    num_paragraphs = extracted_data.get('num_paragraphs', "")
                    lines = extracted_data.get('lines', "")
                    char_count = extracted_data.get('char_count', "")
                    char_count_spaces = extracted_data.get('char_count_spaces', "")
                    word_count = extracted_data.get('word_count', "")
                    total_pages = extracted_data.get('total_pages', "")

                    title = extracted_data.get('title', "")
                    author = extracted_data.get('author', "")
                    creator = extracted_data.get('creator', "")
                    producer = extracted_data.get('producer', "")
                    readable_creation_date = extracted_data.get('readable_creation_date', "")
                    readable_mod_date = extracted_data.get('readable_mod_date', "")

                    
                    
                    auto_generated_summary = auto_generate_summary(extracted_text)
                    #auto_generated_summary =  summarize_and_highlight(extracted_text)
                    summary_percentage = calculate_auto_summary_percentage(auto_generated_summary, text)
                    context['summary_percentage'] = summary_percentage
                    
                

                  
                    auto_generated_summary_word_count = count_words(auto_generated_summary)
                    auto_generated_summary_char_count = character_count(auto_generated_summary)
                    auto_generated_summary_char_count_one = count_char_one(auto_generated_summary)
                    auto_generated_summary_page_count = count_pages(auto_generated_summary)
                    auto_generated_summary_lines = count_lines_helper(auto_generated_summary)
                    auto_generated_summary_para = count_paragraphs_helper(auto_generated_summary)

                    

                    

                    context.update({
                        'extracted_text': extracted_text,
                        'num_paragraphs': num_paragraphs,
                        'lines': lines,
                        'char_count': char_count,
                        'char_count_spaces': char_count_spaces,
                        'word_count': word_count,
                        'auto_generated_summary': auto_generated_summary,
                        'auto_generated_summary_word_count': auto_generated_summary_word_count,
                        'auto_generated_summary_char_count': auto_generated_summary_char_count,
                        'auto_generated_summary_char_count_one': auto_generated_summary_char_count_one,
                        'auto_generated_summary_page_count': auto_generated_summary_page_count,
                        'auto_generated_summary_lines': auto_generated_summary_lines,
                        'auto_generated_summary_para': auto_generated_summary_para
                        
                        
                    })
                    return render(request, 'pdf2.html', context)
                

                

    return render(request, 'pdf2.html', context)









def generate_summary_view(request):
    context = {}

    if request.method == 'POST':
        keywords = request.POST.get("keywords")
        summary_length = request.POST.get("summary_length")

        
        extracted_data = request.session.get('extracted_data', {})
        extracted_text = extracted_data.get('extracted_text', "")
        top_keywords = extracted_data.get('top_keywords', [])
        num_paragraphs = extracted_data.get('num_paragraphs', "")
        lines = extracted_data.get('lines', "")
        char_count = extracted_data.get('char_count', "")
        char_count_spaces = extracted_data.get('char_count_spaces', "")
        word_count = extracted_data.get('word_count', "")
        total_pages = extracted_data.get('total_pages', "")

        title = extracted_data.get('title', "")
        author = extracted_data.get('author', "")
        creator = extracted_data.get('creator', "")
        producer = extracted_data.get('producer', "")
        readable_creation_date = extracted_data.get('readable_creation_date', "")
        readable_mod_date = extracted_data.get('readable_mod_date', "")

        summary = robin(keywords, extracted_text, summary_length)
        summary_percentage = calculate_summary_percentage(summary, extracted_text)
        context['summary_percentage'] = summary_percentage
        


        context.update({
            'context': summary,
            'top_keywords': top_keywords,
            'extracted_text': extracted_text,
            'num_paragraphs': num_paragraphs,
            'lines': lines,
            'char_count': char_count,
            'char_count_spaces': char_count_spaces,
            'word_count': word_count,
            'total_pages': total_pages,

            'title': title,
            'author': author,
            'creator': creator,
            'producer': producer,
            'readable_creation_date': readable_creation_date,
            'readable_mod_date': readable_mod_date,
            'keywords': keywords,


        })

    return render(request, 'pdf2.html', context)







def homepage(request):
    return render(request, 'homepage.html')

def random(request):
    return render(request, 'new.html')


#AUTO GENERATE

# views.py

from django.shortcuts import render
import spacy
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize

nltk.download('punkt')
nltk.download('stopwords')
nlp = spacy.load('en_core_web_sm')
stop_words = set(stopwords.words("english"))

def auto_generate_summary(text):
    
    sentences = sent_tokenize(text)
    # print("sentences",sentences)
    filtered_top_keywords = summarize_text(text, min_keyword_length=3, min_keyword_count=5)
    
  
    tokenized_sentences = [nltk.word_tokenize(sentence.lower()) for sentence in sentences]
    
   
    cleaned_sentences = [
        [word for word in sentence if word not in stop_words]
        for sentence in tokenized_sentences
    ]
    
   
    sentence_scores = [len(sentence) for sentence in cleaned_sentences]
    
    target_length = len(sentences) // 3
    
    
    selected_sentences = sorted(
        range(len(sentences)),
        key=lambda i: sentence_scores[i],
        reverse=True
    )[:target_length]
    
 
    summary = ' '.join(sentences[i] for i in selected_sentences)
    for keyword, _ in filtered_top_keywords:
        # summary = summary.replace(keyword, f"**{keyword}**")
        summary = summary.replace(keyword, f'<span style="font-weight: bold;" class="text-danger">{keyword}</span>')

    
    
    return summary








'''def auto_generate_summary(text, min_keyword_length=3, min_keyword_count=5):
    # Tokenize the input text
    words = word_tokenize(text.lower())
    stop_words = set(stopwords.words('english'))
    punctuation = set(string.punctuation)

    # Filter out stopwords and punctuation
    filtered_words = [word for word in words if word not in stop_words and word not in punctuation]

    # Count keyword occurrences
    keyword_counts = Counter(filtered_words)

    # Get top keywords based on specified criteria
    top_keywords = [(keyword, count) for keyword, count in keyword_counts.items() if len(keyword) >= min_keyword_length and count >= min_keyword_count]

    # Generate summary
    summary = top_keywords

    # Highlight keywords in the summary
    for keyword, _ in top_keywords:
        summary = summary.replace(keyword, f'<span style="background-color: yellow;">{keyword}</span>')

    return summary'''
'''def auto_generate_summary(text, min_keyword_length=3, min_keyword_count=5):
    
    words = word_tokenize(text.lower())
    stop_words = set(stopwords.words('english'))
    punctuation = set(string.punctuation)

    
    filtered_words = [word for word in words if word not in stop_words and word not in punctuation]

    
    keyword_counts = Counter(filtered_words)

    
    top_keywords = [(keyword, count) for keyword, count in keyword_counts.items() if len(keyword) >= min_keyword_length and count >= min_keyword_count]

    
    summary = text.lower()  

    
    for keyword, _ in top_keywords:
        # kirancode
        summary = summary.replace(keyword, f'<span class="text-danger">{keyword}</span>')

    return summary'''
'''import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords

nltk.download('punkt')
nltk.download('stopwords')

def auto_generate_summary(text):
    # Define keywords (add your keywords here)
    keywords = ['keyword1', 'keyword2', 'keyword3']
    
    sentences = sent_tokenize(text)
    stop_words = set(stopwords.words('english'))
        
    tokenized_sentences = [nltk.word_tokenize(sentence.lower()) for sentence in sentences]
    cleaned_sentences = [
        [word for word in sentence if word not in stop_words]
        for sentence in tokenized_sentences
    ]
    sentence_scores = [len(sentence) for sentence in cleaned_sentences]
    
    target_length = len(sentences) // 3
    

    selected_sentences = sorted(
        range(len(sentences)),
        key=lambda i: sentence_scores[i],
        reverse=True
    )[:target_length]

    
    
    summary = ''
    for i in selected_sentences:
        for word in keywords:
            sentences[i] = sentences[i].replace(word, f'<span class="text-danger">{word}</span>')
        summary += sentences[i] + ' '

    return summary'''

'''from docx import Document
from django.http import HttpResponse
import io

def download_summary_as_word(text):
     if request.method == 'post':
        text = request.POST.get('text', '')  # Assuming the form field is named 'text'

        summary = auto_generate_summary(text)

    
        
    

    
    #summary = auto_generate_summary(text)

    
        doc = Document()
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(summary)

    
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

    
        response = HttpResponse(doc_buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=summary.docx'

        return response'''
'''from django.http import HttpResponse
from docx import Document
import io
def download_summary_as_word(request):
    if request.method == 'POST':
        
        text = request.POST.get('summaryText', '') 
        
        
        
        # Assuming the form field is named 'text'

        summary = auto_generate_summary(text)
        
        
        doc = Document()
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(summary)

        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        # print("Document content:", doc_buffer.getvalue())
        response = HttpResponse(summary, content_type='text/plain')
        response['Content-Disposition'] = 'attachment; filename=summary.txt'
        return response'''

        #response = HttpResponse(doc_buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        #response['Content-Disposition'] = 'attachment; filename=summary.docx'
        #return response
from docx import Document
import io
from django.http import HttpResponse

'''def download_summary_as_word(request):
    if request.method == 'POST':
        text = request.POST.get("text","") 

        summary = auto_generate_summary(text)
        
        doc = Document()
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(summary)

        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        

        response = HttpResponse(doc_buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=summary.docx'
        doc.save(response)
        return response'''



'''def download_summary_as_word(text):
    summary = auto_generate_summary(text)
    doc = Document()
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(summary)
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    with open('summary.docx', 'wb') as f:
        f.write(doc_buffer.read())
    return 'summary.docx'
'''
'''def tool(request):
    context = {}
    auto_summary = None
    auto_generated_summary_word_count = None
    auto_generated_summary_char_count = None
    auto_generated_summary_char_count_one = None
    auto_generated_summary_page_count = None
    auto_generated_summary_lines = None
    auto_generated_summary_para = None
    summary_percentage = None
    


    if request.method == 'POST':
        show_statistics = False
        text = request.POST.get("description")
        keywords = request.POST.get("keywords")
        summary_length = request.POST.get("summary_length")

        context['oldDescription'] = text

        if 'extract_keywords' in request.POST:
            cleaned_text = remove_non_ascii(text)
            top_keywords = summarize_text(cleaned_text)
            context['top_keywords'] = top_keywords

        if 'generate_summary' in request.POST:
            cleaned_text = remove_non_ascii(text)
            top_keywords = summarize_text(cleaned_text)
            context['top_keywords'] = top_keywords
            summary = robin(keywords, text, summary_length)
            context['context'] = summary
            context['show_statistics'] = True

            
            summary_percentage = calculate_summary_percentage(summary, text)
            context['summary_percentage'] = summary_percentage

        if 'auto_summary' in request.POST:
            auto_summary =auto_generate_summary(text)
            

            auto_generated_summary_word_count = count_words(auto_summary)
            auto_generated_summary_char_count = character_count(auto_summary)
            auto_generated_summary_char_count_one = count_char_one(auto_summary)
            auto_generated_summary_page_count = count_pages(auto_summary)
            auto_generated_summary_lines = count_lines_helper(auto_summary)
            auto_generated_summary_para = count_paragraphs_helper(auto_summary)
            summary_percentage = calculate_auto_summary_percentage(auto_summary, text)
            context['summary_percentage'] = summary_percentage

        if 'auto_generate_summary' in request.POST:
            auto_summary = auto_generate_summary(text)  # Assuming auto_summary contains the summary text
            
            # Generate DOCX file
            doc = Document()
            doc.add_paragraph(auto_summary)
            
            # In-memory output stream for the document
            f = io.BytesIO()
            doc.save(f)
            length = f.tell()
            f.seek(0)

            response = HttpResponse(f.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename="auto_generated_summary.docx"'
            response['Content-Length'] = length
            return response

    
        count = count_paragraphs_helper(text)
        lines = count_lines_helper(text)
        updated_lines = count_lines(text)
        char_count = character_count(text)
        char_count_one = count_char_one(text)
        word_count = count_words(text)
        pages = count_pages(text)

        show_statistics = True

        context.update({
            'count': count,
            'lines': lines,
            'pages': pages,
            'text': text,
            'updated_lines': updated_lines,
            'char_count': char_count,
            'char_count_one': char_count_one,
            'word_count': word_count,
            'show_statistics': show_statistics,
            'keywords': keywords,
            'auto_generated_summary': auto_summary,
            'auto_generated_summary_word_count': auto_generated_summary_word_count,
            'auto_generated_summary_char_count': auto_generated_summary_char_count,
            'auto_generated_summary_char_count_one': auto_generated_summary_char_count_one,
            'auto_generated_summary_page_count': auto_generated_summary_page_count,
            'auto_generated_summary_lines': auto_generated_summary_lines,
            'auto_generated_summary_para': auto_generated_summary_para,



        })

    return render(request, 'textsum.html', context)'''